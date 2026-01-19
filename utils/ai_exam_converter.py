import json
import re
from docx import Document
from utils.gemini_api import get_gemini_response


def extract_text_from_docx(docx_path):
    """
    Đọc toàn bộ nội dung từ file .docx
    """
    try:
        doc = Document(docx_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
        
        # Đọc cả text trong tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        full_text.append(text)
        
        return '\n'.join(full_text)
    
    except Exception as e:
        raise Exception(f"Lỗi khi đọc file Word: {str(e)}")


def convert_exam_with_ai(docx_text, exam_title="", exam_description=""):
    """
    Sử dụng AI Gemini để chuyển đổi đề thi thành format JSON
    """
    
    # Giới hạn độ dài input để tránh vượt quá token limit
    max_input_length = 15000
    if len(docx_text) > max_input_length:
        docx_text = docx_text[:max_input_length]
        print(f"⚠️ Đã cắt nội dung đề thi xuống {max_input_length} ký tự")
    
    prompt = f"""
Bạn là trợ lý AI chuyên chuyển đổi đề thi. Hãy phân tích nội dung đề thi dưới đây và chuyển thành format JSON.

**QUY TẮC:**
1. Chỉ tạo 2 loại câu hỏi:
   - **Trắc nghiệm ABCD** (type: "tl1"): Có 4 đáp án A, B, C, D - chỉ 1 đáp án đúng
   - **Tự luận** (type: "essay"): Câu hỏi mở, không có đáp án cố định

2. **KHÔNG** tạo câu hỏi Đúng/Sai (type: "tl2")

3. Với câu trắc nghiệm:
   - Phải có đủ 4 đáp án A, B, C, D
   - `correct_answer` là chữ cái (A/B/C/D) - VÍ DỤ: "B" hoặc "C"
   - Nếu đề bài thiếu đáp án, hãy tự tạo thêm đáp án hợp lý

4. Với câu tự luận:
   - Không có options (để object rỗng {{}})
   - `correct_answer` là gợi ý đáp án (nếu có trong đề)
   - Nếu không có gợi ý, để `correct_answer` = ""

5. **GIỮ NỘI DUNG NGẮN GỌN:**
   - `question`: Chỉ lấy nội dung chính, bỏ phần dài dòng
   - `explanation`: Tối đa 1-2 câu, bỏ nếu không cần thiết
   - Nếu đề có >20 câu, chỉ lấy 20 câu đầu tiên

6. Format JSON trả về (KHÔNG có markdown backticks):

{{
  "title": "{exam_title or 'Đề thi'}",
  "description": "{exam_description or 'Đề thi trắc nghiệm'}",
  "time_limit": 15,
  "questions": [
    {{
      "number": 1,
      "question": "Nội dung câu hỏi (ngắn gọn)",
      "type": "tl1",
      "options": {{
        "A": "Đáp án A",
        "B": "Đáp án B",
        "C": "Đáp án C",
        "D": "Đáp án D"
      }},
      "correct_answer": "B",
      "explanation": ""
    }},
    {{
      "number": 2,
      "question": "Câu hỏi tự luận",
      "type": "essay",
      "options": {{}},
      "correct_answer": "",
      "explanation": ""
    }}
  ]
}}

**NỘI DUNG ĐỀ THI:**

{docx_text}

**LƯU Ý QUAN TRỌNG:**
- Nếu đề thi quá dài, chỉ convert tối đa 20 câu đầu tiên
- Trả về ĐÚNG format JSON như trên
- TUYỆT ĐỐI KHÔNG thêm markdown backticks (```) hoặc text giải thích
- CHỈ TRẢ VỀ JSON thuần túy
- Đảm bảo tất cả câu trắc nghiệm có đủ 4 đáp án ABCD
- Phân loại chính xác câu trắc nghiệm (tl1) và tự luận (essay)
"""

    try:
        # Gọi AI với max_tokens cao hơn
        response = get_gemini_response(prompt, temperature=0.5, max_tokens=8192)
        
        # Làm sạch response
        response = response.strip()
        
        # Loại bỏ markdown wrappers
        if response.startswith('```json'):
            response = response[7:]
        elif response.startswith('```'):
            response = response[3:]
        
        if response.endswith('```'):
            response = response[:-3]
        
        response = response.strip()
        
        # Kiểm tra response có bị cắt không
        if not response.endswith('}') and not response.endswith(']'):
            print("⚠️ Response có vẻ bị cắt, đang thử khôi phục...")
            # Tìm JSON object cuối cùng hoàn chỉnh
            last_brace = response.rfind('}')
            if last_brace > 0:
                # Tìm xem có phải đóng array không
                temp = response[:last_brace + 1]
                if temp.count('[') > temp.count(']'):
                    temp += ']'
                if temp.count('{') > temp.count('}'):
                    temp += '}'
                response = temp
            else:
                raise ValueError("Response bị cắt và không thể khôi phục")
        
        # Parse JSON
        try:
            exam_data = json.loads(response)
        except json.JSONDecodeError as e:
            # Log lỗi chi tiết
            error_detail = f"""
❌ JSON Parse Error
Lỗi: {str(e)}
Response length: {len(response)}
First 300 chars: {response[:300]}
Last 300 chars: {response[-300:]}
            """
            print(error_detail)
            raise ValueError(f"AI trả về JSON không hợp lệ: {str(e)}\n\nVui lòng thử lại hoặc chọn đề thi ngắn hơn.")
        
        # Validate
        if 'questions' not in exam_data or not isinstance(exam_data['questions'], list):
            raise ValueError("JSON không có trường 'questions' hợp lệ")
        
        if len(exam_data['questions']) == 0:
            raise ValueError("AI không tạo được câu hỏi nào. Vui lòng kiểm tra nội dung đề thi.")
        
        # Chuẩn hóa dữ liệu
        for idx, question in enumerate(exam_data['questions'], start=1):
            question['number'] = idx
            question['id'] = idx
            
            # Đảm bảo có type
            if 'type' not in question:
                if 'options' in question and question['options']:
                    question['type'] = 'tl1'
                else:
                    question['type'] = 'essay'
            
            # Đảm bảo trắc nghiệm có đủ 4 đáp án
            if question['type'] == 'tl1':
                if 'options' not in question or not isinstance(question['options'], dict):
                    raise ValueError(f"Câu {idx}: Câu trắc nghiệm thiếu đáp án")
                
                # Kiểm tra có đủ 4 đáp án A, B, C, D
                required_keys = {'A', 'B', 'C', 'D'}
                current_keys = set(question['options'].keys())
                
                if not required_keys.issubset(current_keys):
                    missing = required_keys - current_keys
                    raise ValueError(f"Câu {idx}: Thiếu đáp án {', '.join(missing)}")
                
                # Chuẩn hóa correct_answer
                if isinstance(question.get('correct_answer'), list):
                    question['correct_answer'] = question['correct_answer'][0]
                
                if not question.get('correct_answer'):
                    raise ValueError(f"Câu {idx}: Thiếu đáp án đúng")
                
                # Validate đáp án đúng nằm trong A/B/C/D
                correct = str(question['correct_answer']).strip().upper()
                if correct not in required_keys:
                    raise ValueError(f"Câu {idx}: Đáp án đúng '{correct}' không hợp lệ (phải là A, B, C hoặc D)")
                
                question['correct_answer'] = correct
            
            # Tự luận
            elif question['type'] == 'essay':
                question['options'] = {}
                if 'correct_answer' not in question:
                    question['correct_answer'] = ""
            
            # Đảm bảo có explanation
            if 'explanation' not in question:
                question['explanation'] = ""
        
        # Thêm time_limit nếu chưa có
        if 'time_limit' not in exam_data:
            exam_data['time_limit'] = 15
        
        print(f"✅ AI đã tạo {len(exam_data['questions'])} câu hỏi")
        
        return exam_data
    
    except json.JSONDecodeError as e:
        error_msg = f"""
❌ AI trả về JSON không hợp lệ

Lỗi: {str(e)}

Response length: {len(response) if 'response' in locals() else 'N/A'}

First 500 chars:
{response[:500] if 'response' in locals() else 'N/A'}

Last 500 chars:
{response[-500:] if 'response' in locals() else 'N/A'}
        """
        print(error_msg)
        
        raise ValueError(f"AI trả về JSON không hợp lệ. Vui lòng thử lại hoặc chọn đề thi ngắn hơn.")
    
    except Exception as e:
        print(f"❌ Error in convert_exam_with_ai: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"Lỗi khi xử lý với AI: {str(e)}")


def validate_exam_data(exam_data):
    """
    Kiểm tra tính hợp lệ của đề thi sau khi convert
    """
    errors = []
    
    if 'questions' not in exam_data:
        errors.append("Thiếu trường 'questions'")
        return errors
    
    questions = exam_data['questions']
    
    if not questions:
        errors.append("Đề thi không có câu hỏi nào")
        return errors
    
    for idx, q in enumerate(questions, start=1):
        q_type = q.get('type', 'tl1')
        
        if not q.get('question'):
            errors.append(f"Câu {idx}: Thiếu nội dung câu hỏi")
        
        if q_type == 'tl1':
            # Kiểm tra trắc nghiệm
            if 'options' not in q or not isinstance(q['options'], dict):
                errors.append(f"Câu {idx}: Câu trắc nghiệm thiếu đáp án")
                continue
            
            options = q['options']
            if len(options) < 4:
                errors.append(f"Câu {idx}: Thiếu đáp án (cần 4 đáp án ABCD)")
            
            required_keys = {'A', 'B', 'C', 'D'}
            if not required_keys.issubset(set(options.keys())):
                missing = required_keys - set(options.keys())
                errors.append(f"Câu {idx}: Thiếu đáp án {', '.join(missing)}")
            
            correct = q.get('correct_answer')
            if not correct:
                errors.append(f"Câu {idx}: Thiếu đáp án đúng")
            elif str(correct).strip().upper() not in options:
                errors.append(f"Câu {idx}: Đáp án đúng '{correct}' không nằm trong A/B/C/D")
        
        elif q_type == 'essay':
            # Câu tự luận không cần validate nhiều
            pass
        
        else:
            errors.append(f"Câu {idx}: Loại câu hỏi '{q_type}' không hợp lệ (chỉ chấp nhận 'tl1' hoặc 'essay')")
    
    return errors